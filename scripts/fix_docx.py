"""Apply MLSecOps article corrections to Chapter 1.docx (full article)."""
from docx import Document
from copy import deepcopy

DOC_PATH = r"c:\Users\m.haghighian\Desktop\MLSecOps\Chapter 1.docx"

REPLACEMENTS = [
    # Title / abstract
    (
        "MLSecOps is a reference architecture for securing AI-based systems",
        "MLSecOps is a practical security framework for securing AI-based systems",
    ),
    (
        "Each stage in this lifecycle maps to security gates, evidence collection, and controls described in Chapter 6.",
        "Each stage in this lifecycle maps to security gates, evidence collection, and controls described in Chapter 6. The diagram below is an executive view (8 stages); Chapter 6 defines the operational 10-stage pipeline with explicit gates.",
    ),
    # Ch.4 table header glitch
    ("Control\tDescription", "Control\tPurpose"),
    # Ch.5
    (
        "eliminates the need for long-term key management",
        "reduces long-term private key management burden (identity provider trust and verification policy still required)",
    ),
    (
        "a Security Evidence Pack must be stored with the model",
        "an Evidence Pack must be stored with the model",
    ),
    (
        "Proxy gateway for model API keys (BlackVault pattern): agents and pipelines never see raw keys; immediate kill switch on compromise",
        "Proxy gateway for model API keys (API key proxy pattern): agents and pipelines call a gateway that holds credentials; they never receive raw keys, enabling immediate kill switch on compromise",
    ),
    # Ch.6
    (
        "Gates 4, 7, 8, and 9 in the CT cycle are critical and must never be bypassed under any circumstances.",
        "Stages 4, 7, 8, and 9 in the CT cycle are critical and must never be bypassed under any circumstances: Quality Gate 1, Final Security Testing, Final Quality Gate, and Sign Model (see pipeline stage table above—not every numbered stage is a gate).",
    ),
    (
        "Gates 4, 7, 8, and 9 with canary mandatory in every CT.",
        "Mandatory Quality Gate 1, Final Security Testing, Final Quality Gate, and Sign Model in every CT cycle with canary deployment.",
    ),
    (
        "## Alignment with MLOps lifecycle",
        "",  # skip markdown in word
    ),
    (
        "and lintML (ML security linter from Nvidia) for secrets, containers, notebooks, and ML code.",
        "and lintML (ML security linter from Nvidia) for secrets, containers, notebooks, and ML code. Note: lintML runs underlying scanners via Docker containers; CI runners must have Docker available.",
    ),
    # Ch.7 OWASP - individual row replacements
    (
        "Insecure Output Handling",
        "Improper Output Handling (LLM05)",
    ),
    (
        "Overreliance",
        "Misinformation (LLM09; Overreliance removed in OWASP 2025)",
    ),
    (
        "Model Denial of Service",
        "Unbounded Consumption (LLM10)",
    ),
    (
        "OCR/multimodal scan, pattern detection",
        "Unicode normalization, pattern detection, length/entropy heuristics",
    ),
    # Ch.10 MITRE
    ("LLM Safety Bypass", "LLM Jailbreak (AML.T0054)"),
    ("Exfiltration via Inference API", "Exfiltration via AI Inference API (AML.T0024)"),
    ("Poison Retrieval Corpus", "RAG Poisoning (AML.T0070)"),
    ("Poison Agent Memory", "AI Agent Context Poisoning (AML.T0080)"),
    ("Abuse Agent Tooling", "AI Agent Tool Invocation (AML.T0053)"),
    ("Exfiltration of Sensitive Information", "Exfiltration via AI Agent Tool (AML.T0057)"),
    # Ch.11
    (
        "the AI-DAL framework (based on DAL ideas in safety-critical software engineering)",
        "the AI-DAL concept (author-adapted from DAL ideas in safety-critical software engineering; not a published industry standard)",
    ),
    # Ch.12 / Appendix MITRE IDs
    ("AML.T0044", "AML.T0024"),
    ("Poison Web Index", "RAG Poisoning"),
    ("AML.T0066", "AML.T0070"),
    ("Evade ML Model", "Evade AI Model"),
    ("Missing mandatory AI-BOM field", "— (evidence generation; enforce completeness via Conftest/OPA)"),
    (
        "Protect AI, HiddenLayer, Robust Intelligence",
        "HiddenLayer; Protect AI (Palo Alto Networks / Prisma AIRS, 2025); Robust Intelligence (Cisco, 2024)",
    ),
    (
        "# lintML: security linter for ML code (from Nvidia)",
        "# lintML: security linter for ML code (from Nvidia); requires Docker for underlying scanners",
    ),
    # Ch.13
    (
        "LangSmith and API key exposure (2025)",
        "Agent API key exposure pattern (illustrative)",
    ),
    (
        "In some agent scenarios, API keys were extracted via prompt injection or tool chain.",
        "In agent architectures, storing provider API keys in prompts, tool configs, or agent memory creates a realistic exposure path via prompt injection or tool-output manipulation. This is a design pattern to avoid, not a single documented vendor incident.",
    ),
    (
        "In scenarios such as PoisonGPT, an attacker publishes a poisoned model or artifact under a name similar to a legitimate model. A developer may download the wrong model and import it into the pipeline.",
        "In the PoisonGPT demonstration (Mithril Security, 2023), researchers intentionally uploaded a poisoned GPT-2 model to Hugging Face to show that a public registry can deliver a backdoored model that generates attacker-controlled output while appearing legitimate. The risk is supply-chain trust in public model hubs—not name typosquatting alone.",
    ),
    (
        "Control for similar names and typosquatting",
        "Control for similar names and typosquatting (supplementary; PoisonGPT itself was a deliberate poisoned upload, not a naming collision)",
    ),
    # Ch.3 framework bullets (partial matches)
    ("MITRE ATLAS Reconnaissance", "AML.T0067 Discover AI Agent Configuration"),
    ("MITRE ATLAS Persistence", "AML.T0080 AI Agent Context Poisoning"),
    ("MITRE ATLAS Lateral Movement", "AML.T0053 AI Agent Tool Invocation"),
    ("MITRE ATLAS Exfiltration", "AML.T0024 Exfiltration via AI Inference API"),
]

INSERT_AFTER = {
    "## Alignment with MLOps lifecycle": (
        "This guide's 10-stage security pipeline extends the OpenSSF Secure MLOps lifecycle "
        "(9 primary stages in the 2025 whitepaper) by splitting artifact loading, scanning, and "
        "signing into explicit security stages with enforceable gates."
    ),
}


def replace_in_paragraph(paragraph, old, new):
    if old not in paragraph.text:
        return False
    if new == "":
        return False
    # Preserve run formatting imperfectly by full paragraph replace
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Fallback: merge runs
    full = paragraph.text.replace(old, new)
    if full != paragraph.text:
        for i, run in enumerate(paragraph.runs):
            if i == 0:
                run.text = full
            else:
                run.text = ""
        return True
    return False


def process_container(paragraphs):
    count = 0
    for para in paragraphs:
        for old, new in REPLACEMENTS:
            if new and replace_in_paragraph(para, old, new):
                count += 1
    return count


def insert_after_heading(doc, heading_text, insert_text):
    for i, para in enumerate(doc.paragraphs):
        if heading_text.lower() in para.text.lower().replace("\t", " "):
            new_p = doc.paragraphs[i]._element
            # create new paragraph after
            from docx.oxml import OxmlElement
            from docx.text.paragraph import Paragraph

            p = OxmlElement("w:p")
            new_p.addnext(p)
            new_para = Paragraph(p, para._parent)
            new_para.text = insert_text
            return True
    return False


def main():
    doc = Document(DOC_PATH)
    total = process_container(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += process_container(cell.paragraphs)

    # Ch.1 lifecycle mapping note as plain paragraph search
    lifecycle_note = (
        "Executive lifecycle (Chapter 1) maps to pipeline stages (Chapter 6) as follows: "
        "Data Ingest → stages 1–4; Train/Fine-tune → stage 5; Evaluate → stage 6; "
        "Security Test → stage 7; Sign & Register → stages 8–9; Deploy, Runtime Monitor, "
        "and SOC/IR → stage 10 plus Chapter 10."
    )
    for para in doc.paragraphs:
        if (
            "executive view (8 stages)" in para.text
            and lifecycle_note not in para.text
        ):
            para.text = para.text + " " + lifecycle_note
            total += 1
            break

    insert_after_heading(doc, "Alignment with MLOps lifecycle", INSERT_AFTER["## Alignment with MLOps lifecycle"])

    doc.save(DOC_PATH)
    print(f"Applied replacements across docx (approx {total} paragraph/cell updates)")


if __name__ == "__main__":
    main()
