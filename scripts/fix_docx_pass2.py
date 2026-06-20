"""Second pass: fix remaining docx paragraph issues."""
from docx import Document

DOC_PATH = r"c:\Users\m.haghighian\Desktop\MLSecOps\Chapter 1.docx"

LIFECYCLE = (
    "Each stage in this lifecycle maps to security gates, evidence collection, and controls described in Chapter 6. "
    "The diagram below is an executive view (8 stages); Chapter 6 defines the operational 10-stage pipeline with explicit gates. "
    "Executive lifecycle (Chapter 1) maps to pipeline stages (Chapter 6) as follows: Data Ingest → stages 1–4; "
    "Train/Fine-tune → stage 5; Evaluate → stage 6; Security Test → stage 7; Sign & Register → stages 8–9; "
    "Deploy, Runtime Monitor, and SOC/IR → stage 10 plus Chapter 10."
)

POISONGPT = (
    "In the PoisonGPT demonstration (Mithril Security, 2023), researchers intentionally uploaded a poisoned GPT-2 model to "
    "Hugging Face to show that a public registry can deliver a backdoored model that generates attacker-controlled output while "
    "appearing legitimate. The risk is supply-chain trust in public model hubs—not name typosquatting alone."
)

OPENSSF = (
    "This guide's 10-stage security pipeline extends the OpenSSF Secure MLOps lifecycle (9 primary stages in the 2025 whitepaper) "
    "by splitting artifact loading, scanning, and signing into explicit security stages with enforceable gates."
)


def clean(text: str) -> str:
    return text.replace("\xa0", " ")


def set_text(paragraph, text: str):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def main():
    doc = Document(DOC_PATH)
    total = 0

    for para in doc.paragraphs:
        t = clean(para.text)

        if "MLSecOps is a reference architecture for securing AI-based systems" in t:
            set_text(para, t.replace(
                "MLSecOps is a reference architecture for securing AI-based systems",
                "MLSecOps is a practical security framework for securing AI-based systems",
            ))
            total += 1

        if t.strip() == "Each stage in this lifecycle maps to security gates, evidence collection, and controls described in Chapter 6.":
            set_text(para, LIFECYCLE)
            total += 1

        if "In scenarios such as PoisonGPT, an attacker publishes a poisoned model" in t:
            set_text(para, POISONGPT)
            total += 1

        if "the AI-DAL framework (based on DAL ideas in safety-critical software engineering)" in t:
            set_text(para, t.replace(
                "the AI-DAL framework (based on DAL ideas in safety-critical software engineering)",
                "the AI-DAL concept (author-adapted from DAL ideas in safety-critical software engineering; not a published industry standard)",
            ))
            total += 1

        if "lintML (ML security linter from Nvidia) for secrets, containers, notebooks, and ML code." in t and "Docker containers" not in t:
            set_text(para, t.replace(
                "and lintML (ML security linter from Nvidia) for secrets, containers, notebooks, and ML code.",
                "and lintML (ML security linter from Nvidia) for secrets, containers, notebooks, and ML code. "
                "Note: lintML runs underlying scanners via Docker containers; CI runners must have Docker available.",
            ))
            total += 1

    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "Alignment with MLOps lifecycle":
            if i + 1 < len(doc.paragraphs) and OPENSSF not in doc.paragraphs[i + 1].text:
                doc.paragraphs[i + 1].insert_paragraph_before(OPENSSF)
                total += 1
            break

    doc.save(DOC_PATH)
    print(f"Second pass updates: {total}")


if __name__ == "__main__":
    main()
