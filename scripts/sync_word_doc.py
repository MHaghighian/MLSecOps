"""Sync Chapter 1.docx with markdown fixes and MLSecOps Guide v0.1 branding."""
from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

SRC = Path(r"c:\Users\m.haghighian\Desktop\MLSecOps\Chapter 1.docx")
DST = Path(r"c:\Users\m.haghighian\Desktop\MLSecOps\MLSecOps-Guide-v0.1.docx")


def clean(text: str) -> str:
    return text.replace("\xa0", " ").replace("\u2014", "-").replace("\u2013", "-").strip()


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    t = paragraph.text.replace("\xa0", " ")
    if old not in t:
        return False
    set_paragraph_text(paragraph, t.replace(old, new))
    return True


def insert_after(paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = text
    if style:
        new_para.style = style
    return new_para


def replace_in_all(doc: Document, old: str, new: str) -> int:
    count = 0
    for para in doc.paragraphs:
        if replace_in_paragraph(para, old, new):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_in_paragraph(para, old, new):
                        count += 1
    return count


def set_cell(row, col: int, text: str) -> None:
    cell = row.cells[col]
    if cell.paragraphs:
        set_paragraph_text(cell.paragraphs[0], text)
        for p in cell.paragraphs[1:]:
            set_paragraph_text(p, "")
    else:
        cell.text = text


def update_owasp_table(doc: Document) -> None:
    rows_data = [
        (
            "LLM01 Prompt Injection",
            "Bypassing instructions or changing model behavior",
            "Strict System Prompt, Gateway, red team testing",
        ),
        (
            "LLM02 Sensitive Information Disclosure",
            "Disclosure of sensitive or confidential data",
            "DLP, context control, output restrictions",
        ),
        (
            "LLM04 Data and Model Poisoning",
            "Poisoned training data, fine-tuning, or RAG corpus",
            "Data validation, ingest controls, re-index playbook",
        ),
        (
            "LLM05 Improper Output Handling",
            "Unsafe use of model output by another system",
            "Output validation and sandbox",
        ),
        (
            "LLM06 Excessive Agency",
            "Agent or tool actions beyond intended scope",
            "Tool allowlist, Intent Gate, human approval",
        ),
        (
            "LLM08 Vector and Embedding Weaknesses",
            "Poisoned or leaked retrieval/embedding data",
            "ACL at retrieval, tenant isolation, ingest scan",
        ),
        (
            "LLM09 Misinformation",
            "Harmful or unreliable generated content",
            "Human review, grounding, output policy",
        ),
        (
            "LLM10 Unbounded Consumption",
            "High token consumption or expensive requests",
            "Rate limit, quota, and cost monitoring",
        ),
    ]

    for table in doc.tables:
        header = [clean(c.text) for c in table.rows[0].cells]
        if header[:3] == ["Threat", "Description", "Control"] and any(
            "Prompt Injection" in clean(r.cells[0].text) for r in table.rows[1:]
        ):
            # Remove extra rows beyond needed
            while len(table.rows) > len(rows_data) + 1:
                tbl = table._tbl
                tbl.remove(table.rows[-1]._tr)
            # Ensure enough rows
            while len(table.rows) < len(rows_data) + 1:
                table.add_row()
            set_cell(table.rows[0], 0, "Threat (OWASP LLM 2025)")
            for i, row_vals in enumerate(rows_data, start=1):
                for j, val in enumerate(row_vals):
                    set_cell(table.rows[i], j, val)
            return


def update_mitre_table_100(doc: Document) -> None:
    replacements = {
        "Discover AI Assets": "Discover AI Agent Configuration",
        "AML.Txxxx": "",  # handle per row below
    }
    target_rows = [
        ("AI Reconnaissance", "Discover AI Agent Configuration", "AML.T0067"),
        ("Autonomous Agent Abuse", "AI Agent Tool Invocation", "AML.T0053"),
        ("AI Worm Propagation", "AI Agent Context Poisoning", "AML.T0080"),
        ("Model Resource Abuse", "Cost Harvesting", "AML.T0034"),
    ]

    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        header = [clean(c.text) for c in table.rows[0].cells]
        if header != ["Threat", "Technique", "ID"]:
            continue
        flat = " ".join(clean(c.text) for r in table.rows for c in r.cells)
        if "Autonomous Agent Abuse" not in flat:
            continue
        for threat, technique, tid in target_rows:
            for row in table.rows[1:]:
                if clean(row.cells[0].text) == threat:
                    set_cell(row, 1, technique)
                    set_cell(row, 2, tid)
        return


def update_appendix_b_table(doc: Document) -> None:
    extra_rows = [
        ("Jailbreak", "LLM Jailbreak", "AML.T0054"),
        ("Memory Poisoning", "AI Agent Context Poisoning", "AML.T0080"),
        ("Tool Abuse", "AI Agent Tool Invocation", "AML.T0053"),
    ]
    for table in doc.tables:
        header = [clean(c.text) for c in table.rows[0].cells]
        if header != ["Threat", "Technique", "ID"]:
            continue
        flat = " ".join(clean(c.text) for r in table.rows for c in r.cells)
        if "Appendix" in flat or len(table.rows) == 7:
            existing = {clean(r.cells[0].text) for r in table.rows[1:]}
            for threat, technique, tid in extra_rows:
                if threat in existing:
                    for row in table.rows[1:]:
                        if clean(row.cells[0].text) == threat:
                            set_cell(row, 1, technique)
                            set_cell(row, 2, tid)
                else:
                    row = table.add_row()
                    set_cell(row, 0, threat)
                    set_cell(row, 1, technique)
                    set_cell(row, 2, tid)
            # Fix duplicate technique text in Model Extraction cell
            for row in table.rows[1:]:
                if clean(row.cells[0].text) == "Model Extraction":
                    set_cell(row, 1, "Exfiltration via AI Inference API")
                    set_cell(row, 2, "AML.T0024")
            return


def update_ch3_mapping(doc: Document) -> None:
    mapping_lines = [
        "Particular mappings include (technique-level examples, not full coverage):",
        "LLM01 Prompt Injection -> AML.T0051 LLM Prompt Injection",
        "LLM03 Supply Chain -> AML.T0058 Publish Poisoned Models",
        "LLM06 Excessive Agency -> AML.T0053 AI Agent Tool Invocation",
        "LLM08 Vector and Embedding Weaknesses -> AML.T0070 RAG Poisoning",
        "AI reconnaissance -> AML.T0067 Discover AI Agent Configuration",
        "Agent memory/context attacks -> AML.T0080 AI Agent Context Poisoning",
        "Model extraction -> AML.T0024 Exfiltration via AI Inference API",
        "Resource abuse -> AML.T0034 Cost Harvesting",
    ]

    start_idx = None
    for i, para in enumerate(doc.paragraphs):
        if clean(para.text) == "Particular mappings include:":
            start_idx = i
            break
    if start_idx is None:
        return

    set_paragraph_text(doc.paragraphs[start_idx], mapping_lines[0])
    idx = start_idx + 1
    while idx < len(doc.paragraphs):
        t = clean(doc.paragraphs[idx].text)
        if t == "Chapter Summary":
            break
        if t.startswith("LLM") or t.startswith("AML.") or "MITRE ATLAS" in t or "Discover" in t:
            if idx - start_idx < len(mapping_lines) - 1:
                set_paragraph_text(doc.paragraphs[idx], mapping_lines[idx - start_idx])
            else:
                set_paragraph_text(doc.paragraphs[idx], "")
            idx += 1
            continue
        break


def insert_ch3_scope_note(doc: Document) -> None:
    scope = (
        "Scope note: Sections marked emerging describe research-stage or plausible future "
        "capabilities (e.g., autonomous malware at scale, AI worms such as Morris II). "
        "Sections marked demonstrated / active patterns reflect threats with published incidents "
        "or active exploitation patterns (e.g., tool abuse, memory poisoning, compute hijacking). "
        "Threat models should prioritize demonstrated risks first."
    )
    for i, para in enumerate(doc.paragraphs):
        if clean(para.text) == "Chapter 3: Autonomous AI Threats and Offensive AI Operations":
            nxt = doc.paragraphs[i + 1].text if i + 1 < len(doc.paragraphs) else ""
            if "Scope note" in nxt:
                return
            insert_after(para, scope)
            return


def add_owasp_note(doc: Document) -> None:
    note = (
        "Note: Overreliance appeared in OWASP LLM Top 10 (2023) but was removed in the 2025 edition; "
        "related risks are partly covered by LLM09 Misinformation and operational human-review controls."
    )
    for i, para in enumerate(doc.paragraphs):
        if clean(para.text) == "Primary LLM threats":
            # find table after this heading and insert note after table
            j = i + 1
            while j < len(doc.paragraphs) and not doc.paragraphs[j].text.strip():
                j += 1
            # insert before Security controls for LLM
            for k in range(i + 1, min(i + 8, len(doc.paragraphs))):
                if clean(doc.paragraphs[k].text) == "Security controls for LLM":
                    if note not in doc.paragraphs[k - 1].text:
                        doc.paragraphs[k].insert_paragraph_before(note)
                    return


def patch_cover_xml(doc_path: Path) -> None:
    """Normalize cover text boxes in document.xml."""
    with zipfile.ZipFile(doc_path, "r") as zin:
        xml = zin.read("word/document.xml").decode("utf-8")

    # Keep main title; simplify duplicate cover blocks
    xml = xml.replace(
        "A Practical Guide to   Securing AI Systems Across   the Lifecycle",
        "Securing AI Systems Across the Lifecycle",
    )
    xml = re.sub(
        r"(MLSecOps Guide v0\.1\s*)+",
        "MLSecOps Guide v0.1 ",
        xml,
        count=2,
    )

    with zipfile.ZipFile(doc_path, "r") as zin:
        contents = {name: zin.read(name) for name in zin.namelist()}

    contents["word/document.xml"] = xml.encode("utf-8")
    with zipfile.ZipFile(doc_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, data in contents.items():
            zout.writestr(name, data)


def main() -> None:
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    replace_in_all(
        doc,
        "MLSecOps is a practical security guide for securing AI-based systems",
        "MLSecOps Guide (v0.1) is a practical operational guide for securing AI-based systems",
    )
    replace_in_all(
        doc,
        "a Security Evidence Pack must be stored with the model",
        "an Evidence Pack must be stored with the model",
    )
    replace_in_all(
        doc,
        "Proxy gateway for model API keys (BlackVault pattern): agents and pipelines never see raw keys; immediate kill switch on compromise",
        "Proxy gateway for model API keys (API key proxy pattern): agents and pipelines call a gateway that holds credentials; they never receive raw keys, enabling immediate kill switch on compromise",
    )
    replace_in_all(
        doc,
        "MLSecOps is a reference architecture for securing AI-based systems",
        "MLSecOps Guide (v0.1) is a practical operational guide for securing AI-based systems",
    )
    replace_in_all(
        doc,
        "MLSecOps is a practical security framework for securing AI-based systems",
        "MLSecOps Guide (v0.1) is a practical operational guide for securing AI-based systems",
    )

    insert_ch3_scope_note(doc)
    update_ch3_mapping(doc)
    update_owasp_table(doc)
    add_owasp_note(doc)
    update_mitre_table_100(doc)
    update_appendix_b_table(doc)

    doc.save(DST)
    patch_cover_xml(DST)

    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()
